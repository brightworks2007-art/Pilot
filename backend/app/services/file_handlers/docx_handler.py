import io

from docx import Document as DocxDocument

from app.services.file_handlers.base import FileHandler


class DocxHandler(FileHandler):
    """Handles .docx via python-docx."""

    def read_text(self, file_bytes: bytes) -> str:
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)

    def create(self, content_text: str) -> bytes:
        doc = DocxDocument()
        for line in content_text.split("\n"):
            doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def update(self, file_bytes: bytes, instruction_result_text: str) -> bytes:
        # Phase 2 MVP: replace the document body with the new content wholesale.
        # A finer-grained "edit paragraph 3 only" mode can be added once we
        # need it -- python-docx supports targeting individual paragraphs/runs.
        doc = DocxDocument(io.BytesIO(file_bytes))
        for paragraph in list(doc.paragraphs):
            paragraph.text = ""
        for line in instruction_result_text.split("\n"):
            doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
